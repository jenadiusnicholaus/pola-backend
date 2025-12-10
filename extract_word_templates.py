"""
Script to extract Word document templates and convert them to HTML templates
that preserve the original formatting for PDF generation
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json

def extract_word_template(docx_path):
    """Extract content and formatting from a Word document"""
    doc = Document(docx_path)
    
    template_data = {
        'filename': os.path.basename(docx_path),
        'content': [],
        'fields': []  # {{field_name}} placeholders found
    }
    
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
            
        para_data = {
            'text': paragraph.text,
            'alignment': str(paragraph.alignment) if paragraph.alignment else 'LEFT',
            'runs': []
        }
        
        # Extract run-level formatting
        for run in paragraph.runs:
            run_data = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_size': run.font.size.pt if run.font.size else None,
                'font_name': run.font.name,
            }
            para_data['runs'].append(run_data)
            
            # Check for field placeholders like {{field_name}}
            import re
            fields = re.findall(r'\{\{([^}]+)\}\}', run.text)
            template_data['fields'].extend(fields)
        
        template_data['content'].append(para_data)
    
    # Extract tables
    for table in doc.tables:
        table_data = {
            'type': 'table',
            'rows': []
        }
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_data = {
                    'text': cell.text,
                    'paragraphs': []
                }
                for para in cell.paragraphs:
                    cell_data['paragraphs'].append({
                        'text': para.text,
                        'alignment': str(para.alignment) if para.alignment else 'LEFT'
                    })
                row_data.append(cell_data)
            table_data['rows'].append(row_data)
        
        template_data['content'].append(table_data)
    
    return template_data

def convert_to_html_template(template_data):
    """Convert extracted template data to HTML with inline styling"""
    html_parts = []
    html_parts.append('<!DOCTYPE html>')
    html_parts.append('<html>')
    html_parts.append('<head>')
    html_parts.append('<meta charset="UTF-8">')
    html_parts.append('<style>')
    html_parts.append('''
        @page {
            size: A4;
            margin: 2.5cm;
        }
        body {
            font-family: 'Times New Roman', Times, serif;
            font-size: 12pt;
            line-height: 1.8;
            color: #000;
        }
        .center { text-align: center; }
        .left { text-align: left; }
        .right { text-align: right; }
        .justify { text-align: justify; }
        .bold { font-weight: bold; }
        .italic { font-style: italic; }
        .underline { text-decoration: underline; }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 10px 0;
        }
        td, th {
            border: 1px solid #000;
            padding: 5px;
        }
    ''')
    html_parts.append('</style>')
    html_parts.append('</head>')
    html_parts.append('<body>')
    
    for item in template_data['content']:
        if isinstance(item, dict) and item.get('type') == 'table':
            # Render table
            html_parts.append('<table>')
            for row in item['rows']:
                html_parts.append('<tr>')
                for cell in row:
                    html_parts.append('<td>')
                    for para in cell['paragraphs']:
                        alignment_class = 'left'
                        if 'CENTER' in para['alignment']:
                            alignment_class = 'center'
                        elif 'RIGHT' in para['alignment']:
                            alignment_class = 'right'
                        elif 'JUSTIFY' in para['alignment']:
                            alignment_class = 'justify'
                        
                        html_parts.append(f'<p class="{alignment_class}">{para["text"]}</p>')
                    html_parts.append('</td>')
                html_parts.append('</tr>')
            html_parts.append('</table>')
        else:
            # Render paragraph
            alignment_class = 'left'
            if 'CENTER' in item['alignment']:
                alignment_class = 'center'
            elif 'RIGHT' in item['alignment']:
                alignment_class = 'right'
            elif 'JUSTIFY' in item['alignment']:
                alignment_class = 'justify'
            
            html_parts.append(f'<p class="{alignment_class}">')
            
            for run in item['runs']:
                classes = []
                if run['bold']:
                    classes.append('bold')
                if run['italic']:
                    classes.append('italic')
                if run['underline']:
                    classes.append('underline')
                
                style_parts = []
                if run['font_size']:
                    style_parts.append(f'font-size: {run["font_size"]}pt')
                
                class_attr = f' class="{" ".join(classes)}"' if classes else ''
                style_attr = f' style="{"; ".join(style_parts)}"' if style_parts else ''
                
                html_parts.append(f'<span{class_attr}{style_attr}>{run["text"]}</span>')
            
            html_parts.append('</p>')
    
    html_parts.append('</body>')
    html_parts.append('</html>')
    
    return '\n'.join(html_parts)

if __name__ == '__main__':
    sample_templates_dir = '/Users/mac/development/python_projects/pola-backend/sample_templates'
    output_dir = '/Users/mac/development/python_projects/pola-backend/extracted_templates'
    
    os.makedirs(output_dir, exist_ok=True)
    
    # Process each Word document
    for filename in os.listdir(sample_templates_dir):
        if filename.endswith('.docx'):
            print(f'\nProcessing: {filename}')
            docx_path = os.path.join(sample_templates_dir, filename)
            
            try:
                # Extract template data
                template_data = extract_word_template(docx_path)
                
                # Save JSON data
                json_path = os.path.join(output_dir, filename.replace('.docx', '.json'))
                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(template_data, f, indent=2, ensure_ascii=False)
                
                # Convert to HTML
                html_content = convert_to_html_template(template_data)
                html_path = os.path.join(output_dir, filename.replace('.docx', '.html'))
                with open(html_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                
                print(f'  ✓ Extracted {len(template_data["content"])} elements')
                print(f'  ✓ Found {len(set(template_data["fields"]))} unique fields: {set(template_data["fields"])}')
                print(f'  ✓ Saved to: {json_path}')
                print(f'  ✓ HTML template: {html_path}')
                
            except Exception as e:
                print(f'  ✗ Error: {e}')
    
    print(f'\n✓ All templates extracted to: {output_dir}')
